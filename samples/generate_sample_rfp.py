"""
generate_sample_rfp.py — Creates a synthetic RFP DOCX for demo purposes.

Based on the EnABLEHR structure visible in the reference files.
Run once to create samples/sample_rfp.docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REQUIREMENTS = [
    # Feature Group, Title, Description
    ("Create simple customer services and processes",
     "FA-1: Creation of simple customer services and processes",
     "The HR department should be able to create its own, simple customer services and processes digitally. "
     "The resulting products can be made available to employees and fit into the overall architecture of the HR offerings. "
     "It is important to adhere to a uniform presentation and short implementation time."),

    ("Create simple customer services and processes",
     "FA-2: If-then map relationships / decision trees",
     "In order to respond dynamically to user inputs or decisions, HR employees can store simple if-then "
     "relationships in the customer services and processes they have created themselves. No programming knowledge is required."),

    ("Create simple customer services and processes",
     "FA-3: Define activities for person groups",
     "The HR team can store the target group involved in processes for each HR application created by itself. "
     "This enables simple, linear information or approval scenarios to be mapped."),

    ("Create simple customer services and processes",
     "FA-4: Deposit calculation logic for individual data fields",
     "When creating simple customer services, it should be possible to add custom data fields. "
     "These fields can change after input or in response to user decisions. "
     "Programming knowledge is not necessary; a basic understanding of logical relationships must be given."),

    ("Create simple customer services and processes",
     "FA-5: Design simple and intuitive HR application interfaces",
     "Employees can independently create interfaces using modular systems. "
     "They receive direct feedback on their changes and can map simple HR products/services. "
     "There is a fixed set of components (data fields, buttons with fixed actions, icons, etc.)."),

    ("Create simple customer services and processes",
     "FA-6: Different field types for HR data",
     "When designing interfaces, a WYSIWYG editor can be used to select from an existing set of field types "
     "(list fields, text input, date pickers, etc.). HR employees are provided with a suitable range of options."),

    ("Create simple customer services and processes",
     "FA-7: Selection of HR data fields for display and processing",
     "When creating interfaces, it should be possible to use an existing set of HR data fields. "
     "HR employees can quickly create interfaces whose display fields are directly linked to HR databases."),

    ("Create simple customer services and processes",
     "FA-8: Use/creation of design templates and principles",
     "HR employees can rely on templates and design principles to create standardized HR products. "
     "The digital solutions offered to employees should provide consistent user experience."),

    ("Integration of HR applications",
     "FA-9: Integration of HR applications",
     "In order to illustrate end-to-end processes, all HR systems must communicate in a networked manner. "
     "Systems connect via middleware solutions with defined APIs. "
     "A smooth exchange of data and information takes place through API-based integration."),

    ("Integration of HR applications",
     "FA-10: Migrate and transfer HR data",
     "For initial data loading or adding further company participations, HR can work with IT "
     "to execute a defined migration plan in accordance with ETL specifications. "
     "Rapid Deployment Packages should be supported."),

    ("Integration of HR applications",
     "FA-11: Exchange HR data between applications",
     "Individual HR systems must be able to exchange data, information and actions with each other. "
     "Solutions should have modern APIs and interface management options to enable centralised networking."),

    ("Integration of HR applications",
     "FA-12: Establish connectivity between HR applications",
     "Individual solutions are to be implemented using modern connection options "
     "(token-based, JSON interface, PUSH and PULL). "
     "All solutions can be centrally networked in a uniform manner."),

    ("Automation of acceptance of unstructured Customer interactions",
     "FA-13: Automation of unstructured customer interactions (Phone, Mail, Chat)",
     "The volume of unstructured requests processed by HR employees is to be reduced. "
     "HR solutions should interpret and channel customer concerns automatically. "
     "Only cases with consulting requirements involve HR staff directly."),

    ("Automation of acceptance of unstructured Customer interactions",
     "FA-14: Context-sensitive processing and HR process knowledge",
     "In order to provide employees with the best possible support, the HR IT landscape should provide "
     "chatbots, digital assistants, and contextual help. This includes intuitive navigation and process knowledge "
     "on individual HR topics (maternity, sabbatical, etc.)."),

    ("Automation of acceptance of unstructured Customer interactions",
     "FA-15: Context-sensitive processing of employee-related information",
     "To assist employees in entering data or starting processes, only information required for current tasks "
     "should be displayed. Data input only requires information not already available in HR systems."),

    ("Automation of acceptance of unstructured Customer interactions",
     "FA-16: Automation of manual and recurring activities",
     "All repetitive and consistent system tasks currently carried out by HR employees should be automatable. "
     "Activities are carried out by the system itself where no human decisions are necessary."),

    ("Automation of acceptance of unstructured Customer interactions",
     "FA-17: Interpretation of telephone calls",
     "Telephone enquiries made by employees can be interpreted in terms of content. "
     "For selected HR products and processes, their start can be carried out automatically. "
     "Unassigned enquiries are forwarded to the responsible HR staff."),

    ("Automation of acceptance of unstructured Customer interactions",
     "FA-18: Interpretation of text messages (mail, chat)",
     "Unstructured texts sent by employees can be read and interpreted in terms of content. "
     "The start and/or execution of selected HR products and processes can be carried out automatically."),

    ("Ensure the consistency of HR data",
     "FA-19: Ensuring system and application-wide HR data consistency",
     "Uniform data models are provided for mapping end-to-end processes across topics and systems. "
     "Own developments but also SaaS solutions are based on the general data model."),

    ("Ensure the consistency of HR data",
     "FA-20: Single source of truth for HR master data",
     "One authoritative system manages HR master data. "
     "All other systems synchronise from this source via defined APIs. "
     "Data conflicts are detected and flagged automatically."),

    ("Build trust in digital HR systems",
     "FA-21: Build trust through transparency and reliability",
     "Digital HR systems must provide employees with confidence in data accuracy. "
     "Clear audit trails, transparency in decisions, and reliable uptime are mandatory."),

    ("Build trust in digital HR systems",
     "FA-22: Identity and access management",
     "The platform must support Single Sign-On (SSO) and multi-factor authentication. "
     "Role-based access control must be enforced across all HR applications."),

    ("Define target group for access authorization",
     "FA-23: Define target groups for access to HR applications",
     "HR administrators can define target groups and assign access rights to in-house developed HR applications. "
     "Users are assigned roles that control which features and data they can access."),

    ("Development of complex HR applications",
     "FA-24: Development of complex HR applications and customer services",
     "For complex HR use cases that exceed the capabilities of simple services, "
     "a development environment is needed that allows custom applications to be built "
     "as side-by-side extensions without modifying the core HR system."),

    ("Monitoring of resources and applications",
     "FA-25: Monitoring of resources and applications",
     "Standard monitoring capabilities must be available for all deployed HR applications. "
     "Custom monitoring using standard adapters is possible with additional configuration. "
     "Alerting and automated remediation should be supported."),
]

NON_FUNCTIONAL = [
    ("Performance",
     "NFR-1: Response time",
     "All user-facing transactions must complete within 3 seconds under normal load conditions. "
     "Batch processes may take longer but must complete within defined SLA windows."),
    ("Availability",
     "NFR-2: System availability",
     "The HR platform must maintain 99.5% uptime during business hours (06:00–22:00 CET). "
     "Planned maintenance windows must be communicated 5 business days in advance."),
    ("Security",
     "NFR-3: Data protection and GDPR compliance",
     "All personal data must be processed in accordance with GDPR. "
     "Data encryption at rest and in transit is mandatory. "
     "Data residency must be within the EU."),
    ("Scalability",
     "NFR-4: Scalability",
     "The system must support a 50% increase in user load without performance degradation. "
     "Horizontal scaling must be achievable without downtime."),
]


def create_sample_rfp(output_path: str) -> None:
    doc = Document()

    # Title
    title = doc.add_heading("HR IT Transformation — Requirements Specification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Document Reference: EnABLEHR-REQ-001  |  Version: 1.0  |  Status: FINAL")
    doc.add_paragraph()

    # Introduction
    doc.add_heading("1. Introduction", 1)
    doc.add_paragraph(
        "This document specifies the functional and non-functional requirements for the digital transformation "
        "of the HR IT landscape. The objective is to provide HR employees and end users with modern, "
        "integrated, and automated HR services. The solution must be scalable, secure, and aligned with "
        "the organisation's IT architecture principles."
    )

    # IT Architecture Principles
    doc.add_heading("2. IT Architecture Principles", 1)
    principles = [
        "API-first: All integrations must be API-based. No point-to-point connections.",
        "Cloud-first: Prefer cloud-native or cloud-managed services over on-premise deployments.",
        "Security by design: Security controls must be built in from the start, not added later.",
        "Open standards: Use open and widely adopted standards (OAuth 2.0, REST, JSON, OpenAPI).",
        "Modular architecture: Components must be independently deployable and replaceable.",
        "Data sovereignty: Personal data must remain within the EU.",
        "Reusability: Solutions must be designed for reuse across HR use cases.",
    ]
    for p in principles:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_paragraph()

    # Functional Requirements
    doc.add_heading("3. Functional Requirements", 1)

    current_group = None
    for (group, title_text, description) in REQUIREMENTS:
        if group != current_group:
            doc.add_heading(f"3.{list(dict.fromkeys(fg for fg, _, _ in REQUIREMENTS)).index(group) + 1}  {group}", 2)
            current_group = group

        p_title = doc.add_paragraph()
        run = p_title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(description)
        doc.add_paragraph()

    # Non-Functional Requirements
    doc.add_heading("4. Non-Functional Requirements", 1)
    for (group, title_text, description) in NON_FUNCTIONAL:
        p_title = doc.add_paragraph()
        run = p_title.add_run(title_text)
        run.bold = True
        doc.add_paragraph(description)
        doc.add_paragraph()

    doc.save(output_path)
    print(f"Sample RFP created → {output_path}")


if __name__ == "__main__":
    import pathlib
    out = pathlib.Path(__file__).parent / "sample_rfp.docx"
    create_sample_rfp(str(out))
